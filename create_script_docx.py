# -*- coding: utf-8 -*-
# create_script_docx.py — guion basado en xenium_presentacion_final.pptx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\migue\Documents\xenium-spatial-transcriptomics\guion_xenium.docx"

doc = Document()

# ── Margenes ──────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin   = Inches(1.25)
sec.right_margin  = Inches(1.25)

# ── Helpers ───────────────────────────────────────────────────────────────────
BLUE    = RGBColor(0x1F, 0x49, 0x7D)
BROWN   = RGBColor(0x70, 0x50, 0x00)
DARK    = RGBColor(0x20, 0x20, 0x20)
GRAY    = RGBColor(0x70, 0x70, 0x70)

def spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"),  str(after))
    if line:
        sp.set(qn("w:line"),     str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)

def hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    bdr.append(bot)
    pPr.append(bdr)
    spacing(p, 0, 0)

def slide_header(doc, label, title):
    doc.add_paragraph()
    lbl = doc.add_paragraph()
    r   = lbl.add_run(f"[ {label} ]")
    r.font.size = Pt(9); r.font.bold = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    spacing(lbl, 6, 0)
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = BLUE
    spacing(h, 2, 6)

def stage(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"[{text}]")
    r.font.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = BROWN
    spacing(p, 2, 4)

def body(doc, text, bold_phrase=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_phrase and bold_phrase in text:
        i = text.index(bold_phrase)
        for seg, is_bold in [(text[:i], False), (bold_phrase, True), (text[i+len(bold_phrase):], False)]:
            if seg:
                r = p.add_run(seg)
                r.font.size = Pt(11); r.font.bold = is_bold
                r.font.color.rgb = DARK
    else:
        r = p.add_run(text)
        r.font.size = Pt(11); r.font.color.rgb = DARK
    spacing(p, 0, 6, 276)

# ── PORTADA ───────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Guion de Platica")
r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = BLUE
spacing(t, 0, 4)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Panel Xenium: gen + coordenada")
r.font.size = Pt(13); r.font.italic = True
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
spacing(s, 0, 2)

m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run("Transcriptomica espacial de alta resolucion in situ")
r.font.size = Pt(10); r.font.color.rgb = GRAY
spacing(m, 0, 2)

m2 = doc.add_paragraph()
m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m2.add_run("Miguel Angel Diaz Campos  |  Dr. Alfredo de Jesus Rodriguez Gomez  |  Laboratorio de Fallas Medulares  |  2026")
r.font.size = Pt(9); r.font.color.rgb = GRAY
spacing(m2, 0, 8)

info = doc.add_paragraph()
r = info.add_run("Duracion estimada: ")
r.font.bold = True; r.font.size = Pt(10)
r = info.add_run("3:30 - 4:00 min   |   ")
r.font.size = Pt(10)
r = info.add_run("Audiencia: ")
r.font.bold = True; r.font.size = Pt(10)
r = info.add_run("Investigadores sin conocimiento previo de transcriptomica espacial")
r.font.size = Pt(10)
spacing(info, 6, 10)

hr(doc)

# ─────────────────────────────────────────────────────────────────────────────
# DIAPOSITIVA 1 — Por que Xenium
# ─────────────────────────────────────────────────────────────────────────────
slide_header(doc, "DIAPOSITIVA 1 — Por que Xenium",
             "Problema: perdida de ubicacion celular. La solucion: Xenium.")
stage(doc, "Pausa breve al abrir. Mirar al publico.")

body(doc,
    "Cuando queremos entender que genes estan activos en un tejido, la herramienta mas usada "
    "en los ultimos anos ha sido la secuenciacion de RNA de celula unica, el scRNA-seq. Es "
    "poderosa: nos permite medir miles de genes en miles de celulas al mismo tiempo. Pero tiene "
    "un costo importante: para analizar las celulas, primero tenemos que disociar el tejido. "
    "Lo disgregamos, separamos las celulas... y en ese proceso perdemos algo fundamental: "
    "la posicion celular. Sin contexto espacial.",
    bold_phrase="perdemos algo fundamental: la posicion celular.")

body(doc,
    "Y eso importa. Una celula inmune en el centro de un tumor no se comporta igual que una "
    "en el margen. Un hepatocito en zona periportal tiene un programa genico distinto al que "
    "esta en zona centrolobulillar. La funcion de una celula depende de su vecindario.",
    bold_phrase="La funcion de una celula depende de su vecindario.")

body(doc,
    "Aqui entra Xenium. Con Xenium medimos alrededor de cinco mil genes directamente en el tejido "
    "intacto, sin disgregarlo. Cada transcripto conserva sus coordenadas X e Y originales. "
    "Ademas, con el Xenium Protein Bundle 2025, podemos medir proteinas en el mismo corte de "
    "tejido, sin procesamiento adicional, usando anticuerpos conjugados. RNA y proteinas, "
    "al mismo tiempo, en el mismo tejido.",
    bold_phrase="Xenium")

hr(doc)

# ─────────────────────────────────────────────────────────────────────────────
# DIAPOSITIVA 2 — Tecnologia
# ─────────────────────────────────────────────────────────────────────────────
slide_header(doc, "DIAPOSITIVA 2 — Tecnologia",
             "Descifrando el tejido: RNA y proteinas molecula a molecula")
stage(doc, "Avanzar diapositiva.")

body(doc, "Pero, como funciona esto en la practica?")

body(doc,
    "Tomamos una seccion de tejido, tipicamente fijada en formalina, el tipo de muestra estandar "
    "en patologia, y aplicamos dos tipos de reactivos: sondas tipo padlock para el RNA, y "
    "anticuerpos conjugados para proteinas. Estas moleculas se unen especificamente a sus blancos "
    "dentro del tejido intacto, sin necesidad de extraer ni mover nada.",
    bold_phrase="sondas tipo padlock")

body(doc,
    "Una vez hibridadas, las sondas se amplifican mediante amplificacion en circulo rodante, "
    "un proceso que genera pequenas nanobolas de ADN fluorescentes. Cada nanobola es visible al "
    "microscopio y representa un transcripto en una posicion exacta del tejido.",
    bold_phrase="amplificacion en circulo rodante")

body(doc,
    "Despues, el equipo realiza multiples ciclos de imagen usando cuatro colores fluorescentes "
    "distintos en cada ronda. La combinacion de senales a lo largo de todos los ciclos genera "
    "un codigo de barras unico para cada gen o proteina. El resultado: millones de transcritos "
    "y proteinas, todos con coordenadas X e Y precisas. Un mapa molecular de alta resolucion "
    "del tejido.",
    bold_phrase="codigo de barras unico")

hr(doc)

# ─────────────────────────────────────────────────────────────────────────────
# DIAPOSITIVA 3 — Analisis computacional
# ─────────────────────────────────────────────────────────────────────────────
slide_header(doc, "DIAPOSITIVA 3 — Analisis computacional",
             "Deteccion de identidades celulares")
stage(doc, "Avanzar diapositiva.")

body(doc, "Con ese mapa en mano, comienza el analisis computacional.")

body(doc,
    "El primer paso es la segmentacion celular: usando las imagenes del nucleo con DAPI y de "
    "la membrana, definimos los limites de cada celula y asignamos los transcritos y proteinas "
    "que caen dentro de ella. Asi construimos una matriz de expresion genica por celula, "
    "exactamente como en scRNA-seq, pero ahora cada celula sigue teniendo sus coordenadas "
    "en el tejido.",
    bold_phrase="segmentacion celular")

body(doc,
    "A esa matriz le aplicamos herramientas clasicas de bioinformatica: reduccion de dimensiones "
    "con PCA y UMAP, y clustering para identificar grupos de celulas con perfiles de expresion "
    "similares. Luego anotamos esos grupos usando marcadores especificos del tejido que "
    "estemos estudiando, ya sea higado, pulmon, o cualquier otro.")

body(doc,
    "Y aqui viene la ventaja diferencial de Xenium: los analisis espaciales. Podemos preguntar "
    "quien vive junto a quien, si ciertos tipos celulares se concentran en zonas especificas, "
    "o como cambia la composicion del tejido a lo largo de un gradiente. Cuantificamos "
    "vecindarios celulares y sus estadisticas.",
    bold_phrase="analisis espaciales")

body(doc,
    "En conclusion: Xenium integra imagen, expresion genica y proteinas en un solo experimento, "
    "con un pipeline reproducible de codigo abierto. Pasamos del pixel a la identidad celular, "
    "y de la identidad celular a la arquitectura del tejido.",
    bold_phrase="Xenium integra imagen, expresion genica y proteinas en un solo experimento")

stage(doc, "Pausa para preguntas.")

hr(doc)

# ── Tabla de tiempos ──────────────────────────────────────────────────────────
doc.add_paragraph()
h = doc.add_paragraph()
r = h.add_run("Tiempos de referencia")
r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLUE
spacing(h, 4, 6)

from docx.oxml import OxmlElement as OE

table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
table.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ["Seccion", "Tiempo aprox."]
rows_data = [
    ("Diapositiva 1 — Por que Xenium",           "~1:20 min"),
    ("Diapositiva 2 — Tecnologia",               "~1:10 min"),
    ("Diapositiva 3 — Analisis computacional",   "~1:10 min"),
    ("Total",                                    "~3:40 min"),
]

for i, hdr in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.paragraphs[0].clear()
    r    = cell.paragraphs[0].add_run(hdr)
    r.font.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OE("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1F497D")
    tcPr.append(shd)

for r_idx, (col1, col2) in enumerate(rows_data):
    row  = table.rows[r_idx + 1]
    fill = "DCE6F1" if r_idx % 2 == 0 else "FFFFFF"
    bold = (col1 == "Total")
    for c_idx, text in enumerate([col1, col2]):
        cell = row.cells[c_idx]
        cell.paragraphs[0].clear()
        run  = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10); run.font.bold = bold
        if bold: run.font.color.rgb = BLUE
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OE("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill)
        tcPr.append(shd)

# ── Notas de delivery ─────────────────────────────────────────────────────────
doc.add_paragraph()
h = doc.add_paragraph()
r = h.add_run("Notas de delivery")
r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLUE
spacing(h, 10, 4)

notes = [
    ("Diapositiva 1, columna comparativa",
     "senalar visualmente la columna izquierda (celulas disgregadas) y la derecha (tejido intacto) al mencionar el contraste."),
    ("Frase clave",
     '"La funcion de una celula depende de su vecindario" -- hacer pausa breve, es el gancho central.'),
    ("Diapositiva 2, diagrama de flujo",
     "senalar el flujo Tejido -> RNA/Proteina -> Ciclos imagen -> Mapa molecular al explicar los pasos."),
    ("Diapositiva 3, cierre",
     'El ultimo parrafo retoma el titulo "Panel Xenium: gen + coordenada" -- cierra el hilo narrativo de la plática.'),
]

for bold_part, normal_part in notes:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(bold_part + ": ")
    r.font.bold = True; r.font.size = Pt(10)
    r = p.add_run(normal_part)
    r.font.size = Pt(10)
    spacing(p, 0, 4)

doc.save(OUTPUT)
print(f"Guardado: {OUTPUT}")
